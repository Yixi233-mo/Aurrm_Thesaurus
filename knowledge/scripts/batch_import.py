#!/usr/bin/env python3
"""
金融知识库批量导入脚本

功能：将金融数据目录中的所有文档批量导入到知识库系统。
流程：
    1. 扫描数据目录，识别所有待导入文件
    2. 按文件类型分发处理路径
       - PDF → pdf_to_md → md_img → document_split → item_name → bge_embedding → milvus → kg
       - MD  → md_img → document_split → item_name → bge_embedding → milvus → kg
       - DOCX → 文本提取为 MD 后走 MD 分支
    3. 每份文件完成后记录结果
    4. 输出汇总报告

用法：
    cd E:\work_space\掌柜智库\002\knowledge
    python scripts/batch_import.py

环境要求：
    - 已配置 .env（Milvus / Neo4j / OpenAI / MinerU）
    - 网络可访问 MinerU API 和 LLM API
    - Milvus / Neo4j 服务已启动
"""

import os
import sys
import json
import time
import shutil
import logging
import traceback
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple

# ==================== 路径设置 ====================

# 项目根目录（当前文件位于 knowledge/scripts/，往上两级到 knowledge/）
SCRIPT_DIR = Path(__file__).resolve().parent
KNOWLEDGE_DIR = SCRIPT_DIR.parent
PROJECT_DIR = KNOWLEDGE_DIR.parent
sys.path.insert(0, str(PROJECT_DIR))

# 数据目录
DATA_DIR = Path(r"E:\work_space\掌柜智库\002\金融\数据")

# 日志目录
LOG_DIR = KNOWLEDGE_DIR / "logs"
LOG_DIR.mkdir(exist_ok=True)

# ==================== 提前加载 .env ====================
# 必须在使用任何 knowledge 模块之前加载，确保所有子模块都能读取到环境变量
from knowledge.core import env  # noqa: E402 - 加载项目根目录 .env（基于 __file__ 定位，不受 CWD 影响）

# ==================== 日志配置 ====================


def setup_logging() -> logging.Logger:
    """配置日志系统，同时输出到控制台和文件"""
    logger = logging.getLogger("batch_import")
    logger.setLevel(logging.DEBUG)

    fmt = logging.Formatter(
        "%(asctime)s | %(levelname)-5s | %(message)s",
        datefmt="%H:%M:%S"
    )

    # 控制台 handler
    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.INFO)
    ch.setFormatter(fmt)
    logger.addHandler(ch)

    # 文件 handler
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    fh = logging.FileHandler(LOG_DIR / f"batch_import_{ts}.log", encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)
    logger.addHandler(fh)

    return logger


logger = setup_logging()

# ==================== 环境检查 ====================


def check_env() -> bool:
    """检查关键环境变量和依赖"""
    from dotenv import load_dotenv
    load_dotenv(KNOWLEDGE_DIR / ".env")

    checks = {
        "OPENAI_API_KEY": os.getenv("OPENAI_API_KEY"),
        "OPENAI_API_BASE": os.getenv("OPENAI_API_BASE"),
        "MILVUS_URL": os.getenv("MILVUS_URL"),
        "CHUNKS_COLLECTION": os.getenv("CHUNKS_COLLECTION"),
        "NEO4J_URI": os.getenv("NEO4J_URI"),
        "NEO4J_USERNAME": os.getenv("NEO4J_USERNAME"),
        "NEO4J_PASSWORD": os.getenv("NEO4J_PASSWORD"),
        "MINERU_KEY": os.getenv("MINERU_KEY"),
        "BGE_M3_PATH": os.getenv("BGE_M3_PATH"),
    }

    missing = [k for k, v in checks.items() if not v]
    if missing:
        logger.error(f"缺少必要环境变量: {missing}")
        return False

    # 检查 BGE 模型路径是否存在
    bge_path = Path(checks["BGE_M3_PATH"])
    if not bge_path.exists():
        logger.error(f"BGE-M3 模型路径不存在: {bge_path}")
        return False

    logger.info("环境变量检查通过")
    return True


def check_services() -> bool:
    """检查外部服务可用性"""
    import requests

    # 检查 MinerU
    mineru_key = os.getenv("MINERU_KEY", "")
    if mineru_key:
        try:
            resp = requests.get(
                "https://mineru.net/api/v4/extract-results/batch/test",
                headers={"Authorization": f"Bearer {mineru_key}"},
                timeout=10,
            )
            logger.info(f"MinerU 服务状态: HTTP {resp.status_code}")
        except Exception as e:
            logger.warning(f"MinerU 连接检查失败: {e}")

    # 检查 Milvus
    try:
        from pymilvus import MilvusClient
        milvus_url = os.getenv("MILVUS_URL", "")
        if milvus_url:
            client = MilvusClient(uri=milvus_url)
            logger.info("Milvus 连接正常")
        else:
            logger.warning("MILVUS_URL 未配置，跳过 Milvus 检查")
    except Exception as e:
        logger.warning(f"Milvus 连接检查失败: {e}")

    # 检查 Neo4j
    try:
        from knowledge.tools.neo4j_utils import get_neo4j_driver
        driver = get_neo4j_driver()
        with driver.session() as session:
            session.run("RETURN 1")
        logger.info("Neo4j 连接正常")
    except Exception as e:
        logger.warning(f"Neo4j 连接检查失败: {e}")

    return True

# ==================== DOCX 转换 ====================


def convert_docx_to_md(docx_path: Path, output_path: Path) -> bool:
    """将 .docx 文件转换为 Markdown 格式

    使用 python-docx 提取文本，按段落组织为 Markdown。
    """
    try:
        import docx

        doc = docx.Document(str(docx_path))
        md_lines = []

        # 提取标题
        title = docx_path.stem
        md_lines.append(f"# {title}")
        md_lines.append("")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                md_lines.append("")
                continue

            # 根据样式判断标题级别
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name or "标题 1" in style_name:
                md_lines.append(f"\n## {text}\n")
            elif "Heading 2" in style_name or "标题 2" in style_name:
                md_lines.append(f"\n### {text}\n")
            elif "Heading 3" in style_name or "标题 3" in style_name:
                md_lines.append(f"\n#### {text}\n")
            elif text.startswith("Q:") or text.startswith("问：") or text.startswith("问:"):
                md_lines.append(f"\n**{text}**\n")
            elif text.startswith("A:") or text.startswith("答：") or text.startswith("答:"):
                md_lines.append(f"{text}\n")
            else:
                md_lines.append(text)

        # 提取表格
        for table in doc.tables:
            md_lines.append("")
            # 表头
            header = "| " + " | ".join(
                cell.text.strip().replace("|", "\\|") for cell in table.rows[0].cells
            ) + " |"
            md_lines.append(header)
            # 分隔线
            md_lines.append("| " + " | ".join("---" for _ in table.rows[0].cells) + " |")
            # 数据行
            for row in table.rows[1:]:
                line = "| " + " | ".join(
                    cell.text.strip().replace("|", "\\|") for cell in row.cells
                ) + " |"
                md_lines.append(line)
            md_lines.append("")

        md_content = "\n".join(md_lines)
        output_path.write_text(md_content, encoding="utf-8")
        logger.info(f"DOCX 转换成功: {docx_path.name} → {output_path.name}")
        return True

    except ImportError:
        logger.warning("未安装 python-docx，尝试安装...")
        import subprocess
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "python-docx", "-q"]
            )
            # 重试
            return convert_docx_to_md(docx_path, output_path)
        except Exception as e:
            logger.error(f"安装 python-docx 失败: {e}")
            return False
    except Exception as e:
        logger.error(f"DOCX 转换失败 {docx_path}: {e}")
        return False


def convert_doc_to_pdf(doc_path: Path, output_path: Path) -> bool:
    """将 .doc 文件转换为 PDF

    在 Windows 上使用 COM 自动化调用 Word。
    """
    try:
        import comtypes.client  # type: ignore

        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(doc_path))
        doc.SaveAs(str(output_path), FileFormat=17)  # 17 = PDF
        doc.Close()
        word.Quit()
        logger.info(f"DOC → PDF 转换成功: {doc_path.name}")
        return True

    except ImportError:
        logger.warning("未安装 comtypes，尝试安装...")
        import subprocess
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "comtypes", "-q"]
            )
            return convert_doc_to_pdf(doc_path, output_path)
        except Exception as e:
            logger.error(f"安装 comtypes 失败: {e}")
            return False
    except Exception as e:
        logger.error(f"DOC 转换失败 {doc_path}: {e}")
        # 清理可能残留的 Word 进程
        try:
            import subprocess
            subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"], capture_output=True)
        except Exception:
            pass
        return False


def convert_to_importable(input_path: Path, work_dir: Path) -> Optional[Path]:
    """将任意输入文件转换为导入流程可处理的格式（PDF 或 MD）

    Returns:
        转换后的文件路径（PDF 或 MD），转换失败返回 None
    """
    suffix = input_path.suffix.lower()

    if suffix == ".pdf":
        return input_path

    elif suffix == ".md":
        # 复制到工作目录
        dest = work_dir / input_path.name
        shutil.copy2(input_path, dest)
        return dest

    elif suffix == ".docx":
        md_path = work_dir / (input_path.stem + ".md")
        if convert_docx_to_md(input_path, md_path):
            return md_path
        return None

    elif suffix == ".doc":
        pdf_path = work_dir / (input_path.stem + ".pdf")
        if convert_doc_to_pdf(input_path, pdf_path):
            return pdf_path
        return None

    else:
        logger.warning(f"不支持的文件类型: {suffix} ({input_path.name})")
        return None

# ==================== 导入流程执行 ====================


def run_single_import(
    file_path: Path,
    category: str,
    work_dir: Path,
    timeout: int = 600,
) -> Dict:
    """对单个文件执行完整导入流程

    Args:
        file_path: 源文件路径
        category: 文件分类（用于日志）
        work_dir: 工作目录
        timeout: 单文件超时时间（秒）

    Returns:
        结果字典: {
            "file": 文件名,
            "category": 分类,
            "status": "success" | "failed" | "skipped",
            "chunks": 切片数,
            "item_name": 识别出的产品名,
            "error": 错误信息（如有）,
            "duration": 耗时（秒）,
        }
    """
    result = {
        "file": file_path.name,
        "category": category,
        "status": "failed",
        "chunks": 0,
        "item_name": "",
        "error": "",
        "duration": 0,
    }

    start_time = time.time()

    try:
        # Step 1: 转换文件格式
        logger.info(f"[{category}] 开始处理: {file_path.name}")
        converted_path = convert_to_importable(file_path, work_dir)
        if converted_path is None:
            result["error"] = "文件格式转换失败"
            result["status"] = "skipped"
            return result

        # Step 2: 构建初始状态
        import_suffix = converted_path.suffix.lower()

        initial_state = {
            "task_id": f"batch_{int(time.time())}_{file_path.stem[:20]}",
            "is_pdf_read_enabled": (import_suffix == ".pdf"),
            "is_md_read_enabled": (import_suffix == ".md"),
            "file_dir": str(work_dir),
            "import_file_path": str(converted_path),
            "pdf_path": str(converted_path) if import_suffix == ".pdf" else "",
            "md_path": str(converted_path) if import_suffix == ".md" else "",
            "file_title": converted_path.stem,
            "md_content": "",
            "chunks": [],
            "item_name": "",
        }

        # Step 3: 如果是 MD 文件，先读取内容
        if import_suffix == ".md":
            content = converted_path.read_text(encoding="utf-8")
            initial_state["md_content"] = content

        # Step 4: 运行导入图
        from knowledge.processor.import_process.main_graph import kb_import_app

        final_state = None
        for event in kb_import_app.stream(initial_state):
            for key, value in event.items():
                if key != "__end__":
                    logger.debug(f"  [{file_path.name}] 节点完成: {key}")
                final_state = value

        if final_state is None:
            final_state = initial_state

        # Step 5: 收集结果
        chunks = final_state.get("chunks", [])
        item_name = final_state.get("item_name", "")

        result["chunks"] = len(chunks)
        result["item_name"] = item_name
        result["status"] = "success"
        logger.info(
            f"[{category}] ✓ {file_path.name} → "
            f"{len(chunks)} 个切片, 产品名: {item_name}"
        )

    except Exception as e:
        result["error"] = f"{type(e).__name__}: {e}"
        logger.error(
            f"[{category}] ✗ {file_path.name} 失败: {e}\n"
            f"{traceback.format_exc()}"
        )

    finally:
        result["duration"] = round(time.time() - start_time, 1)

    return result

# ==================== 数据扫描 ====================


def scan_data_files(data_dir: Path) -> Dict[str, List[Path]]:
    """扫描数据目录，按子目录分类收集文件

    Returns:
        {"分类名": [文件路径列表], ...}
    """
    categories: Dict[str, List[Path]] = {}

    if not data_dir.exists():
        logger.error(f"数据目录不存在: {data_dir}")
        return categories

    for subdir in sorted(data_dir.iterdir()):
        if not subdir.is_dir():
            continue

        category_name = subdir.name
        files: List[Path] = []

        for f in sorted(subdir.rglob("*")):
            if f.is_file() and not f.name.startswith("."):
                files.append(f)

        if files:
            categories[category_name] = files
            logger.info(f"发现分类 [{category_name}]: {len(files)} 个文件")
        else:
            logger.warning(f"分类 [{category_name}] 中没有找到文件")

    return categories

# ==================== 汇总报告 ====================


def generate_report(results: List[Dict], report_path: Path):
    """生成 Markdown 格式的导入报告"""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 统计
    total = len(results)
    success = sum(1 for r in results if r["status"] == "success")
    failed = sum(1 for r in results if r["status"] == "failed")
    skipped = sum(1 for r in results if r["status"] == "skipped")
    total_chunks = sum(r["chunks"] for r in results)
    total_duration = sum(r["duration"] for r in results)

    lines = [
        f"# 金融知识库导入报告",
        f"",
        f"**生成时间**: {ts}",
        f"",
        f"## 汇总统计",
        f"",
        f"| 指标 | 数值 |",
        f"|------|------|",
        f"| 文件总数 | {total} |",
        f"| 成功 | {success} |",
        f"| 失败 | {failed} |",
        f"| 跳过 | {skipped} |",
        f"| 总切片数 | {total_chunks} |",
        f"| 总耗时 | {total_duration:.1f} 秒 |",
        f"",
        f"## 详细结果",
        f"",
        f"| 分类 | 文件名 | 状态 | 切片数 | 产品名称 | 耗时(s) | 错误 |",
        f"|------|--------|------|--------|----------|---------|------|",
    ]

    for r in results:
        status_icon = {"success": "✅", "failed": "❌", "skipped": "⏭️"}.get(
            r["status"], "?"
        )
        error_short = (r["error"] or "")[:50]
        lines.append(
            f"| {r['category']} | {r['file']} | {status_icon} {r['status']} | "
            f"{r['chunks']} | {r['item_name'] or '-'} | {r['duration']} | "
            f"{error_short or '-'} |"
        )

    lines += ["", "---", ""]

    report_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info(f"报告已保存: {report_path}")

    # 同时输出 JSON 格式结果
    json_path = report_path.with_suffix(".json")
    json_path.write_text(
        json.dumps(
            {
                "timestamp": ts,
                "summary": {
                    "total": total,
                    "success": success,
                    "failed": failed,
                    "skipped": skipped,
                    "total_chunks": total_chunks,
                    "total_duration": total_duration,
                },
                "results": results,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    logger.info(f"JSON 结果已保存: {json_path}")

# ==================== 主流程 ====================


def main():
    """主入口"""
    logger.info("=" * 60)
    logger.info("金融知识库批量导入工具")
    logger.info(f"数据目录: {DATA_DIR}")
    logger.info(f"项目目录: {KNOWLEDGE_DIR}")
    logger.info("=" * 60)

    # 1. 环境检查
    if not check_env():
        logger.error("环境检查失败，请检查 .env 配置")
        sys.exit(1)

    # 2. 服务可用性检查
    check_services()

    # 3. 扫描数据文件
    categories = scan_data_files(DATA_DIR)
    if not categories:
        logger.error("未找到任何数据文件，请检查数据目录")
        sys.exit(1)

    total_files = sum(len(files) for files in categories.values())
    logger.info(f"共发现 {total_files} 个待导入文件")

    # 4. 创建工作目录
    batch_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    work_dir = Path(os.getenv("MD_ROOT_DIR", "./temp-files/")) / f"batch_{batch_id}"
    work_dir.mkdir(parents=True, exist_ok=True)
    logger.info(f"工作目录: {work_dir}")

    # 5. 逐个导入
    results: List[Dict] = []

    for category, files in categories.items():
        logger.info(f"\n{'='*40}")
        logger.info(f"处理分类: {category} ({len(files)} 个文件)")
        logger.info(f"{'='*40}")

        for file_path in files:
            # 为每个文件创建独立子目录
            file_work_dir = work_dir / file_path.stem
            file_work_dir.mkdir(exist_ok=True)

            result = run_single_import(file_path, category, file_work_dir)
            results.append(result)

            # 文件间短暂休息，避免并发压力
            time.sleep(1)

    # 6. 生成报告
    report_path = LOG_DIR / f"import_report_{batch_id}.md"
    generate_report(results, report_path)

    # 7. 终端输出摘要
    logger.info("\n" + "=" * 60)
    logger.info("导入完成！")
    logger.info("=" * 60)

    success = [r for r in results if r["status"] == "success"]
    failed = [r for r in results if r["status"] == "failed"]
    skipped = [r for r in results if r["status"] == "skipped"]

    logger.info(f"总计: {len(results)} 个文件")
    logger.info(f"  成功: {len(success)} 个")
    logger.info(f"  失败: {len(failed)} 个")
    logger.info(f"  跳过: {len(skipped)} 个")
    logger.info(f"  总切片: {sum(r['chunks'] for r in results)} 个")

    if failed:
        logger.info("\n失败文件:")
        for r in failed:
            logger.info(f"  ❌ [{r['category']}] {r['file']}: {r['error']}")

    if skipped:
        logger.info("\n跳过文件:")
        for r in skipped:
            logger.info(f"  ⏭️  [{r['category']}] {r['file']}: {r['error']}")

    logger.info(f"\n详细报告: {report_path}")


if __name__ == "__main__":
    main()
